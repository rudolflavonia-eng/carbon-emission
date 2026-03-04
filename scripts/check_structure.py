"""Check paragraphs 11-70 to find chapter headings and TOC"""
from docx import Document
doc = Document(r'E:\2026next\碳排放\基于Django的碳排放分析及预测可视化平台的设计与实现.docx')

print('=== PARAGRAPHS 11-70 ===')
for i in range(11, min(70, len(doc.paragraphs))):
    text = doc.paragraphs[i].text.strip()
    style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else 'None'
    if text:
        print(f'{i}: [{style}] "{text[:70]}"')
    else:
        print(f'{i}: [{style}] (empty)')

# Also check paragraphs around 50
print('\n=== PARAGRAPHS 44-55 ===')
for i in range(44, min(55, len(doc.paragraphs))):
    text = doc.paragraphs[i].text.strip()
    style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else 'None'
    print(f'{i}: [{style}] "{text[:80]}"')

# Check paragraphs around h2 "5.3 可视化大屏展示" 
print('\n=== PARAGRAPHS 547-575 (end) ===')
for i in range(547, min(575, len(doc.paragraphs))):
    text = doc.paragraphs[i].text.strip()
    style = doc.paragraphs[i].style.name if doc.paragraphs[i].style else 'None'
    if text:
        print(f'{i}: [{style}] "{text[:80]}"')
