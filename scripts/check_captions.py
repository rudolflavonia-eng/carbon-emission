"""Check paragraphs around img_ref and look for figure/table captions"""
import pickle, re
from docx import Document

doc = Document(r'E:\2026next\碳排放\基于Django的碳排放分析及预测可视化平台的设计与实现.docx')

with open(r'E:\2026next\碳排放\scripts\classifications.pkl', 'rb') as f:
    classifications = pickle.load(f)

# Build index
cls_map = {i: cls for i, cls, text in classifications}

# Show context around img_ref
print('=== CONTEXT AROUND IMAGE REFERENCES ===')
for i, cls, text in classifications:
    if cls == 'img_ref':
        # Show 2 paragraphs before and after
        for j in range(max(0,i-2), min(len(doc.paragraphs), i+3)):
            p = doc.paragraphs[j]
            tag = cls_map.get(j, '???')
            print(f'  {j}: [{tag}] "{p.text.strip()[:80]}"')
        print()

# Also search for any paragraph containing 图X-Y or 表X-Y patterns
print('\n=== PARAGRAPHS CONTAINING 图X-Y OR 表X-Y patterns ===')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.search(r'图\d+-\d+', text) and len(text) < 80:
        tag = cls_map.get(i, '???')
        print(f'  {i}: [{tag}] "{text}"')
    if re.search(r'表\d+-\d+', text) and len(text) < 80:
        tag = cls_map.get(i, '???')
        print(f'  {i}: [{tag}] "{text}"')

# Check all headings
print('\n=== ALL HEADINGS ===')
for i, cls, text in classifications:
    if cls in ('h1', 'h2', 'h3'):
        print(f'  {i}: [{cls}] "{text}"')
