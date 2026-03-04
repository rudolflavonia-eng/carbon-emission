"""Verify the rebuilt docx: check headings, captions, figure numbering, code blocks"""
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r'E:\2026next\碳排放\基于Django的碳排放分析及预测可视化平台的设计与实现_排版后.docx')

print("=== Headings ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading'):
        text = p.text.strip()[:60]
        print(f"  [{i}] {p.style.name}: {text}")

print("\n=== Figure/Table Captions ===")
fig_nums = []
tab_nums = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    m = re.match(r'(图\d+-\d+)\s', text)
    if m:
        fig_nums.append(m.group(1))
        print(f"  [{i}] FIG: {text[:50]}")
    m = re.match(r'(表\d+-\d+)\s', text)
    if m:
        tab_nums.append(m.group(1))
        print(f"  [{i}] TAB: {text[:50]}")

print(f"\nFigure count: {len(fig_nums)}, Table count: {len(tab_nums)}")
print(f"Figures: {', '.join(fig_nums)}")
print(f"Tables: {', '.join(tab_nums)}")

print("\n=== Code Blocks ===")
code_count = 0
for p in doc.paragraphs:
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None and pPr.find(qn('w:shd')) is not None:
        code_count += 1
print(f"Code paragraphs: {code_count}")

print("\n=== Tables ===")
for ti, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell = table.rows[0].cells[0].text[:30] if rows > 0 else ''
    print(f"  Table {ti+1}: {rows} rows x {cols} cols, first cell: '{first_cell}'")

# Count Chinese chars in body text
cn_count = 0
for p in doc.paragraphs:
    pPr = p._element.find(qn('w:pPr'))
    has_shd = pPr is not None and pPr.find(qn('w:shd')) is not None
    if has_shd:
        continue  # skip code
    cn_count += len(re.findall(r'[\u4e00-\u9fff]', p.text))
print(f"\nChinese chars (excl code): {cn_count}")
