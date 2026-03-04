"""Round 3: Analyze headers, footers, abstract, and references"""
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

doc = Document(r'E:\2026next\碳排放\论文1.docx')

# Headers and footers
print('=== HEADERS & FOOTERS ===')
for i, sec in enumerate(doc.sections):
    print(f'\nSection {i}:')
    # Header
    header = sec.header
    if header and not header.is_linked_to_previous:
        for p in header.paragraphs:
            text = p.text.strip()
            if text:
                style = p.style.name if p.style else 'None'
                align = p.alignment
                run_info = []
                for r in p.runs[:2]:
                    rPr = r._element.find(qn('w:rPr'))
                    sz = cn = None
                    if rPr is not None:
                        szEl = rPr.find(qn('w:sz'))
                        if szEl is not None: sz = szEl.get(qn('w:val'))
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None: cn = rFonts.get(qn('w:eastAsia'))
                    run_info.append(f'sz={sz} cn={cn}')
                print(f'  Header: [{style}] align={align} text="{text}" runs={run_info}')
                # Check border
                pPr = p._element.find(qn('w:pPr'))
                if pPr is not None:
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        bottom = pBdr.find(qn('w:bottom'))
                        if bottom is not None:
                            val = bottom.get(qn('w:val'))
                            sz_b = bottom.get(qn('w:sz'))
                            print(f'    Border bottom: val={val} sz={sz_b}')
    else:
        print('  Header: linked to previous')
    
    # Footer
    footer = sec.footer
    if footer and not footer.is_linked_to_previous:
        for p in footer.paragraphs:
            text = p.text.strip()
            style = p.style.name if p.style else 'None'
            align = p.alignment
            # Check for PAGE field
            has_page = False
            for r in p.runs:
                if 'PAGE' in r.text:
                    has_page = True
            # Check XML for fldChar
            xml_str = etree.tostring(p._element, encoding='unicode')
            if 'PAGE' in xml_str or 'fldChar' in xml_str:
                has_page = True
            print(f'  Footer: [{style}] align={align} text="{text}" has_page_field={has_page}')
    else:
        print('  Footer: linked to previous')

# Find abstract and reference sections
print('\n=== KEY SECTIONS ===')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text in ['摘要', '摘  要', 'Abstract', 'ABSTRACT', '参考文献', '致谢', '致  谢']:
        style = para.style.name
        align = para.alignment
        pPr = para._element.find(qn('w:pPr'))
        spacing_info = ''
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                spacing_info = f'line={spacing.get(qn("w:line"))} before={spacing.get(qn("w:before"))} after={spacing.get(qn("w:after"))}'
        
        run_info = []
        for r in para.runs[:2]:
            rPr = r._element.find(qn('w:rPr'))
            sz = cn = bold = None
            if rPr is not None:
                szEl = rPr.find(qn('w:sz'))
                if szEl is not None: sz = szEl.get(qn('w:val'))
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None: cn = rFonts.get(qn('w:eastAsia'))
                bEl = rPr.find(qn('w:b'))
                if bEl is not None:
                    bval = bEl.get(qn('w:val'))
                    bold = bval != '0' if bval else True
            run_info.append(f'sz={sz} cn={cn} bold={bold}')
        print(f'Para {i}: "{text}" [{style}] align={align} {spacing_info} runs={run_info}')

# Abstract paragraphs (30-60)
print('\n=== ABSTRACT REGION (para 30-60) ===')
for i in range(30, min(60, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()[:80]
    if not text: continue
    style = para.style.name
    align = para.alignment
    run_info = []
    for r in para.runs[:2]:
        rPr = r._element.find(qn('w:rPr'))
        sz = cn = bold = None
        if rPr is not None:
            szEl = rPr.find(qn('w:sz'))
            if szEl is not None: sz = szEl.get(qn('w:val'))
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None: cn = rFonts.get(qn('w:eastAsia'))
            bEl = rPr.find(qn('w:b'))
            if bEl is not None:
                bval = bEl.get(qn('w:val'))
                bold = bval != '0' if bval else True
        run_info.append(f'sz={sz} cn={cn} bold={bold}')
    print(f'  {i}: [{style}] align={align} "{text}" runs={run_info}')

# Reference section
print('\n=== REFERENCE SECTION ===')
in_ref = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text == '参考文献':
        in_ref = True
    if in_ref:
        if not text and i > 290: continue
        style = para.style.name
        align = para.alignment
        pPr = para._element.find(qn('w:pPr'))
        ind_info = ''
        if pPr is not None:
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                ind_info = f'fl={ind.get(qn("w:firstLine"))} hang={ind.get(qn("w:hanging"))} left={ind.get(qn("w:left"))}'
        run_info = []
        for r in para.runs[:1]:
            rPr = r._element.find(qn('w:rPr'))
            sz = cn = en = None
            if rPr is not None:
                szEl = rPr.find(qn('w:sz'))
                if szEl is not None: sz = szEl.get(qn('w:val'))
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    cn = rFonts.get(qn('w:eastAsia'))
                    en = rFonts.get(qn('w:ascii'))
            run_info.append(f'sz={sz} cn={cn} en={en}')
        print(f'  {i}: [{style}] align={align} "{text[:60]}" {ind_info} runs={run_info}')
        if text == '致  谢' or text == '致谢':
            break
