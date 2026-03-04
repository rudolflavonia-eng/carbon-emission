"""
rebuild_docx.py  –  Convert updated 论文.md → formatted docx
Uses thesis_utils.py for low-level formatting.
Template specs (from analysis):
  Normal:  宋体/TNR 小四(12pt/sz=24) 两端对齐 固定值20磅行距 首行缩进2字符
  H1:      黑体 三号(16pt/sz=32) 居中 段前1行段后1行 keepNext/keepLines pageBreakBefore
  H2:      宋体/TNR 小三(14pt/sz=28) 加粗 左对齐 段前1行段后1行
  H3:      宋体/TNR 小四(12pt) 加粗 段前0.5行段后0.5行
  Caption: 楷体 五号(10.5pt/sz=21) 居中
  Code:    Consolas/宋体 五号(10.5pt) 灰底 左对齐 无缩进
  Ref:     宋体/TNR 五号(10.5pt) 悬挂缩进420
"""
import re, os, sys
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from thesis_utils import (
    set_run_font, fmt_all_runs, set_paragraph_fmt,
    clear_indent, add_code_border, set_hanging_indent,
    add_page_break_before, set_keep_next, ensure_ppr
)

# ── helpers ──────────────────────────────────────────────
def add_run(para, text, cn='宋体', en='Times New Roman', sz=12, b=False):
    run = para.add_run(text)
    set_run_font(run, cn=cn, en=en, sz=sz, b=b)
    return run

def fmt_normal(para, indent=True):
    """Apply Normal formatting"""
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      line_spacing=20, line_rule='exact')
    if indent:
        set_paragraph_fmt(para, first_indent_chars=200)
    fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=False)

def fmt_h1(para):
    para.style = doc.styles['Heading 1']
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line_spacing=20, line_rule='exact',
                      before_lines=1, after_lines=1)
    clear_indent(para)
    fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=16, b=False)
    set_keep_next(para)
    add_page_break_before(para)

def fmt_h2(para):
    para.style = doc.styles['Heading 2']
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                      line_spacing=20, line_rule='exact',
                      before_lines=1, after_lines=1)
    clear_indent(para)
    fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=14, b=True)
    set_keep_next(para)

def fmt_h3(para):
    para.style = doc.styles['Heading 3']
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                      line_spacing=20, line_rule='exact',
                      before_lines=0.5, after_lines=0.5)
    clear_indent(para)
    fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=True)
    set_keep_next(para)

def fmt_caption(para):
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER,
                      line_spacing=20, line_rule='exact')
    clear_indent(para)
    fmt_all_runs(para, cn='楷体', en='Times New Roman', sz=10.5, b=False)

def fmt_code(para):
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT,
                      line_spacing=20, line_rule='exact')
    clear_indent(para)
    add_code_border(para)
    fmt_all_runs(para, cn='宋体', en='Consolas', sz=10.5, b=False)

def fmt_ref(para):
    set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      line_spacing=20, line_rule='exact')
    clear_indent(para)
    set_hanging_indent(para, hanging=420, left=420)
    fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=10.5, b=False)

def add_empty_line(para_before=None):
    """Add empty paragraph after last content"""
    p = doc.add_paragraph('')
    set_paragraph_fmt(p, line_spacing=20, line_rule='exact')
    return p

# ── Parse markdown ─────────────────────────────────────
md_path = r'E:\2026next\碳排放\论文.md'
with open(md_path, 'r', encoding='utf-8') as f:
    md_text = f.read()

lines = md_text.split('\n')

# ── Build elements list ─────────────────────────────────
elements = []  # each: (type, content)  or (type, rows_list) for tables
i = 0
in_code = False
code_buf = []

while i < len(lines):
    line = lines[i]
    
    # Code block toggle
    if line.strip().startswith('```'):
        if not in_code:
            in_code = True
            code_buf = []
            i += 1
            continue
        else:
            # End code block
            in_code = False
            elements.append(('code_block', code_buf[:]))
            code_buf = []
            i += 1
            continue
    
    if in_code:
        code_buf.append(line)
        i += 1
        continue
    
    stripped = line.strip()
    
    # Skip empty lines
    if not stripped:
        i += 1
        continue
    
    # Horizontal rule
    if stripped == '---':
        i += 1
        continue
    
    # H1: # Title  or  ## 第X章
    if re.match(r'^#\s+', line) and not line.startswith('##'):
        text = re.sub(r'^#\s+', '', line).strip()
        elements.append(('title', text))
        i += 1
        continue
    
    # 第X章 headings  (## 第X章)
    m = re.match(r'^##\s+(第\d+章\s+.+)', line)
    if m:
        elements.append(('h1', m.group(1)))
        i += 1
        continue
    
    # Special headings: 摘要, Abstract, 参考文献, 致谢, 附录
    m = re.match(r'^##\s+(.+)', line)
    if m:
        text = m.group(1).strip()
        if text in ['摘  要', '摘要', 'Abstract', '参考文献', '致  谢', '致谢'] or text.startswith('附'):
            elements.append(('h1_special', text))
            i += 1
            continue
    
    # H2: ### X.Y
    m = re.match(r'^###\s+(\d+\.\d+\s+.+)', line)
    if m:
        elements.append(('h2', m.group(1)))
        i += 1
        continue
    
    # H3: #### X.Y.Z
    m = re.match(r'^####\s+(\d+\.\d+\.\d+\s+.+)', line)
    if m:
        elements.append(('h3', m.group(1)))
        i += 1
        continue
    
    # H3-like subsection: #### 附录X
    m = re.match(r'^####\s+(.+)', line)
    if m:
        elements.append(('h3_special', m.group(1)))
        i += 1
        continue
    
    # Sub-heading bold: **（X）...**
    m = re.match(r'^\*\*（\d+）.+\*\*$', stripped)
    if m:
        text = stripped.replace('**', '')
        elements.append(('sub_heading', text))
        i += 1
        continue
    
    # Image reference: ![alt](path)
    m = re.match(r'^!\[(.+?)\]\((.+?)\)', stripped)
    if m:
        elements.append(('image', (m.group(1), m.group(2))))
        i += 1
        continue
    
    # Caption: <center>图X-Y ... </center> or <center>表X-Y ...</center>
    m = re.match(r'^<center>(.+?)</center>$', stripped)
    if m:
        elements.append(('caption', m.group(1)))
        i += 1
        continue
    
    # Table: starts with |
    if stripped.startswith('|'):
        table_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            row_text = lines[i].strip()
            cells = [c.strip() for c in row_text.split('|')[1:-1]]
            # Skip separator rows like |:---:|
            if cells and all(re.match(r'^:?-+:?$', c) for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
        elements.append(('table', table_rows))
        continue
    
    # Reference item: [N] ...
    m = re.match(r'^\[(\d+)\]\s+', stripped)
    if m:
        elements.append(('ref', stripped))
        i += 1
        continue
    
    # Keywords bold: **关键词：** or **Keywords:**
    if stripped.startswith('**关键词') or stripped.startswith('**Keywords'):
        text = stripped.replace('**', '')
        elements.append(('keywords', text))
        i += 1
        continue
    
    # Normal paragraph
    elements.append(('body', stripped))
    i += 1

print(f"Parsed {len(elements)} elements from markdown")

# ── Create Document ──────────────────────────────────────
doc = Document()

# ── Configure styles ─────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for sname in ['Heading 1', 'Heading 2', 'Heading 3']:
    s = doc.styles[sname]
    s.font.name = 'Times New Roman'

# ── Page setup ───────────────────────────────────────────
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

# ── Render elements ──────────────────────────────────────
first_h1 = True  # don't page break on first H1 (title)

for idx, (etype, content) in enumerate(elements):
    
    if etype == 'title':
        # Main title
        p = doc.add_paragraph()
        add_run(p, content, cn='黑体', en='Times New Roman', sz=16, b=True)
        set_paragraph_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                          line_spacing=20, line_rule='exact',
                          before_lines=1, after_lines=1)
        clear_indent(p)
    
    elif etype == 'h1':
        p = doc.add_paragraph()
        add_run(p, content, cn='黑体', en='Times New Roman', sz=16, b=False)
        fmt_h1(p)
    
    elif etype == 'h1_special':
        p = doc.add_paragraph()
        add_run(p, content, cn='黑体', en='Times New Roman', sz=16, b=False)
        fmt_h1(p)
    
    elif etype == 'h2':
        p = doc.add_paragraph()
        add_run(p, content, cn='宋体', en='Times New Roman', sz=14, b=True)
        fmt_h2(p)
    
    elif etype == 'h3':
        p = doc.add_paragraph()
        add_run(p, content, cn='宋体', en='Times New Roman', sz=12, b=True)
        fmt_h3(p)
    
    elif etype == 'h3_special':
        p = doc.add_paragraph()
        add_run(p, content, cn='宋体', en='Times New Roman', sz=12, b=True)
        fmt_h3(p)
    
    elif etype == 'sub_heading':
        p = doc.add_paragraph()
        add_run(p, content, cn='宋体', en='Times New Roman', sz=12, b=True)
        set_paragraph_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line_spacing=20, line_rule='exact',
                          first_indent_chars=200)
    
    elif etype == 'body':
        p = doc.add_paragraph()
        add_run(p, content)
        fmt_normal(p, indent=True)
    
    elif etype == 'keywords':
        p = doc.add_paragraph()
        # Split into bold label and normal text
        if '：' in content:
            label, text = content.split('：', 1)
            add_run(p, label + '：', cn='宋体', en='Times New Roman', sz=12, b=True)
            add_run(p, text, cn='宋体', en='Times New Roman', sz=12, b=False)
        else:
            add_run(p, content, cn='宋体', en='Times New Roman', sz=12, b=False)
        set_paragraph_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line_spacing=20, line_rule='exact',
                          first_indent_chars=200)
    
    elif etype == 'caption':
        # Insert a line break (empty para) before caption
        blank = doc.add_paragraph('')
        set_paragraph_fmt(blank, line_spacing=20, line_rule='exact')
        
        p = doc.add_paragraph()
        add_run(p, content, cn='楷体', en='Times New Roman', sz=10.5, b=False)
        fmt_caption(p)
    
    elif etype == 'image':
        alt, path = content
        p = doc.add_paragraph()
        # Try to insert image
        full_path = os.path.join(r'E:\2026next\碳排放', path.replace('%20', ' '))
        if os.path.exists(full_path):
            run = p.add_run()
            run.add_picture(full_path, width=Cm(14))
        else:
            # Placeholder
            add_run(p, f'[{alt}]', cn='楷体', en='Times New Roman', sz=10.5)
        set_paragraph_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                          line_spacing=20, line_rule='exact')
        clear_indent(p)
    
    elif etype == 'code_block':
        for code_line in content:
            p = doc.add_paragraph()
            add_run(p, code_line if code_line else ' ', cn='宋体', en='Consolas', sz=10.5)
            fmt_code(p)
    
    elif etype == 'table':
        if not content:
            continue
        # Create table
        nrows = len(content)
        ncols = len(content[0]) if content else 0
        if ncols == 0:
            continue
        
        table = doc.add_table(rows=nrows, cols=ncols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Apply "Table Grid" style
        try:
            table.style = doc.styles['Table Grid']
        except:
            pass
        
        for ri, row_cells in enumerate(content):
            for ci, cell_text in enumerate(row_cells):
                if ci < len(table.rows[ri].cells):
                    cell = table.rows[ri].cells[ci]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    add_run(p, cell_text, cn='宋体', en='Times New Roman', sz=10.5)
                    set_paragraph_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER,
                                      line_spacing=20, line_rule='exact')
                    clear_indent(p)
    
    elif etype == 'ref':
        p = doc.add_paragraph()
        add_run(p, content, cn='宋体', en='Times New Roman', sz=10.5)
        fmt_ref(p)

# ── Save ─────────────────────────────────────────────────
out_path = r'E:\2026next\碳排放\基于Django的碳排放分析及预测可视化平台的设计与实现_排版后.docx'
doc.save(out_path)
print(f"Saved to {out_path}")

# ── Verify ───────────────────────────────────────────────
doc2 = Document(out_path)
total_paras = len(doc2.paragraphs)
total_tables = len(doc2.tables)
print(f"Total paragraphs: {total_paras}, tables: {total_tables}")

# Count headings
h1_count = sum(1 for p in doc2.paragraphs if p.style.name == 'Heading 1')
h2_count = sum(1 for p in doc2.paragraphs if p.style.name == 'Heading 2')
h3_count = sum(1 for p in doc2.paragraphs if p.style.name == 'Heading 3')
print(f"H1: {h1_count}, H2: {h2_count}, H3: {h3_count}")
