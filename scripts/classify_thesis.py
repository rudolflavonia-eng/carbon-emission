"""Classify paragraphs in the thesis document"""
import re, pickle
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'E:\2026next\碳排放\基于Django的碳排放分析及预测可视化平台的设计与实现.docx')

classifications = []
in_code_block = False
in_ref = False
in_ack = False
abs_cn_idx = None
abs_en_idx = None
kw_cn_idx = None
kw_en_idx = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else 'Normal'
    
    # Empty
    if not text:
        classifications.append((i, 'empty', text))
        continue
    
    # Code block markers
    if text.startswith('```'):
        in_code_block = not in_code_block
        classifications.append((i, 'code_marker', text))
        continue
    
    if in_code_block:
        classifications.append((i, 'code', text))
        continue
    
    # First paragraph = title
    if i == 0:
        classifications.append((i, 'paper_title', text))
        continue
    
    # Abstract titles
    if text in ['摘要', '摘  要']:
        abs_cn_idx = i
        classifications.append((i, 'abs_cn_title', text))
        continue
    if text in ['Abstract', 'ABSTRACT']:
        abs_en_idx = i
        classifications.append((i, 'abs_en_title', text))
        continue
    
    # Keywords
    if text.startswith('关键词') or text.startswith('关 键 词'):
        kw_cn_idx = i
        classifications.append((i, 'kw_cn', text))
        continue
    if text.lower().startswith('key') and ('words' in text.lower() or 'word' in text.lower()):
        kw_en_idx = i
        classifications.append((i, 'kw_en', text))
        continue
    
    # Reference title
    if text == '参考文献':
        in_ref = True
        classifications.append((i, 'ref_title', text))
        continue
    
    # Acknowledgement title
    if text in ['致谢', '致  谢']:
        in_ref = False
        in_ack = True
        classifications.append((i, 'ack_title', text))
        continue
    
    # Reference items: [number] 
    if in_ref and re.match(r'^\[\d+\]', text):
        classifications.append((i, 'ref_item', text))
        continue
    
    # Acknowledgement body
    if in_ack:
        classifications.append((i, 'ack_body', text))
        continue
    
    # Headings by style
    if style == 'Heading 1' or (re.match(r'^\d+\s{2,}', text) and len(text) < 30):
        classifications.append((i, 'h1', text))
        continue
    if style == 'Heading 2' or (re.match(r'^\d+\.\d+\s{2,}', text) and len(text) < 40):
        classifications.append((i, 'h2', text))
        continue
    if style == 'Heading 3' or (re.match(r'^\d+\.\d+\.\d+\s{2,}', text) and len(text) < 40):
        classifications.append((i, 'h3', text))
        continue
    
    # Also check style-based heading detection
    if style == 'Heading 1':
        classifications.append((i, 'h1', text))
        continue
    if style == 'Heading 2':
        classifications.append((i, 'h2', text))
        continue
    if style == 'Heading 3':
        classifications.append((i, 'h3', text))
        continue
    
    # Figure captions
    if re.match(r'^图\d+-\d+\s', text) or re.match(r'^图\d+-\d+$', text):
        classifications.append((i, 'fig', text))
        continue
    
    # Table captions
    if re.match(r'^表\d+-\d+\s', text) or re.match(r'^表\d+-\d+$', text):
        classifications.append((i, 'tab', text))
        continue
    
    # Image ref
    if text.startswith('!['):
        classifications.append((i, 'img_ref', text))
        continue
    
    # Default body
    classifications.append((i, 'body', text))

# Second pass: refine abstract body
refined = []
for idx, (i, cls, text) in enumerate(classifications):
    if cls == 'body':
        # Check if between abs_cn_title and kw_cn
        if abs_cn_idx is not None and kw_cn_idx is not None:
            if i > abs_cn_idx and i < kw_cn_idx:
                refined.append((i, 'abs_cn', text))
                continue
        # Check if between abs_en_title and kw_en
        if abs_en_idx is not None and kw_en_idx is not None:
            if i > abs_en_idx and i < kw_en_idx:
                refined.append((i, 'abs_en', text))
                continue
    refined.append((i, cls, text))

classifications = refined

# Statistics
from collections import Counter
stats = Counter(cls for _, cls, _ in classifications)
print('=== CLASSIFICATION STATISTICS ===')
for cls, count in sorted(stats.items()):
    print(f'  {cls}: {count}')

print(f'\nTotal paragraphs: {len(doc.paragraphs)}')
print(f'Total classified: {len(classifications)}')

# Print all non-body, non-empty
print('\n=== NON-BODY PARAGRAPHS ===')
for i, cls, text in classifications:
    if cls not in ('body', 'empty', 'abs_cn', 'abs_en', 'ack_body'):
        print(f'  {i}: [{cls}] "{text[:60]}"')

# Save pickle
with open(r'E:\2026next\碳排放\scripts\classifications.pkl', 'wb') as f:
    pickle.dump(classifications, f)
print('\nClassifications saved to pickle.')

# Print figure/table captions for numbering check
print('\n=== FIGURE CAPTIONS ===')
for i, cls, text in classifications:
    if cls == 'fig':
        print(f'  Para {i}: "{text}"')

print('\n=== TABLE CAPTIONS ===')
for i, cls, text in classifications:
    if cls == 'tab':
        print(f'  Para {i}: "{text}"')
