"""Analyze current thesis content: word count per section"""
import re, pickle
from docx import Document

doc = Document(r'E:\2026next\碳排放\基于Django的碳排放分析及预测可视化平台的设计与实现_排版后.docx')

# Count Chinese characters
def count_cn(text):
    return len(re.findall(r'[\u4e00-\u9fff]', text))

def count_all(text):
    """Count Chinese chars + English words (approx)"""
    cn = len(re.findall(r'[\u4e00-\u9fff]', text))
    # English words count as ~2 chars each
    en_words = len(re.findall(r'[a-zA-Z]+', text))
    return cn + en_words * 2

# Collect text by section
current_chapter = "前置"
sections = {}
sections[current_chapter] = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name
    
    if not text:
        continue
    
    # Detect chapter change
    m = re.match(r'^第(\d+)章', text)
    if m:
        current_chapter = f"第{m.group(1)}章 {text}"
        sections[current_chapter] = []
        continue
    
    if text == '参考文献':
        current_chapter = '参考文献'
        sections[current_chapter] = []
        continue
    if text in ['致  谢', '致谢']:
        current_chapter = '致谢'
        sections[current_chapter] = []
        continue
    if text in ['附  录', '附录']:
        current_chapter = '附录'
        sections[current_chapter] = []
        continue
    
    sections.setdefault(current_chapter, []).append(text)

total_cn = 0
total_all = 0
print('=== WORD COUNT BY SECTION ===')
for sec, texts in sections.items():
    full = '\n'.join(texts)
    cn = count_cn(full)
    all_c = count_all(full)
    total_cn += cn
    total_all += all_c
    # Count code vs text
    code_chars = 0
    text_chars = 0
    for t in texts:
        if any(kw in t for kw in ['def ', 'class ', 'import ', 'return ', 'function', 'var ', 'const ', '<div', '<nav', '<svg', '.then(', 'filter(', '.objects']):
            code_chars += len(t)
        else:
            text_chars += count_cn(t)
    print(f'{sec[:40]:<42} 汉字={cn:>5}  估算字数={all_c:>5}  纯文本汉字={text_chars:>5}')

print(f'\n总汉字数: {total_cn}')
print(f'总估算字数(汉字+英文词*2): {total_all}')

# Also count paragraphs with actual body text (not code, not headings)
body_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    style = para.style.name
    if not text:
        continue
    if style in ['Heading 1', 'Heading 2', 'Heading 3']:
        continue
    if style.startswith('toc'):
        continue
    # Skip code (has shading)
    from docx.oxml.ns import qn
    pPr = para._element.find(qn('w:pPr'))
    has_shd = False
    if pPr is not None:
        has_shd = pPr.find(qn('w:shd')) is not None
    if has_shd:
        continue
    # Skip very short empty-ish lines
    if len(text) < 3:
        continue
    # Skip references
    if re.match(r'^\[\d+\]', text):
        continue
    # Skip figure/table captions
    if re.match(r'^图\d+-\d+', text) or re.match(r'^表\d+-\d+', text):
        continue
    # Skip image refs
    if text.startswith('!['):
        continue
    
    body_text.append(text)

pure_body = '\n'.join(body_text)
pure_cn = count_cn(pure_body)
pure_all = count_all(pure_body)
print(f'\n纯正文（排除代码、标题、目录、参考文献、图表注）:')
print(f'  汉字数: {pure_cn}')
print(f'  估算字数: {pure_all}')
print(f'  段落数: {len(body_text)}')

# Show heading structure for expansion planning
print('\n=== CHAPTER STRUCTURE ===')
for para in doc.paragraphs:
    style = para.style.name
    text = para.text.strip()
    if style == 'Heading 1' and text:
        print(f'\n{text}')
    elif style == 'Heading 2' and text:
        print(f'  {text}')
    elif style == 'Heading 3' and text:
        print(f'    {text}')
