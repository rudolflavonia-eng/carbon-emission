"""Quick verify: check which paragraphs got page breaks and heading styles"""
from docx import Document
from docx.oxml.ns import qn
import re

doc = Document(r'E:\2026next\碳排放\基于Django的碳排放分析及预测可视化平台的设计与实现_排版后.docx')

print('=== Heading 1 paragraphs ===')
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 1':
        text = para.text.strip()[:60]
        # Check page break
        pPr = para._element.find(qn('w:pPr'))
        has_pb = False
        if pPr is not None:
            pb = pPr.find(qn('w:pageBreakBefore'))
            has_pb = pb is not None
        print(f'  {i}: "{text}" pageBreak={has_pb}')

print('\n=== Code block check (sample) ===')
code_found = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    pPr = para._element.find(qn('w:pPr'))
    has_shd = False
    if pPr is not None:
        shd = pPr.find(qn('w:shd'))
        has_shd = shd is not None
    if has_shd and text:
        code_found += 1
        if code_found <= 3:
            run = para.runs[0] if para.runs else None
            if run:
                rPr = run._element.find(qn('w:rPr'))
                cn = en = sz = None
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        cn = rFonts.get(qn('w:eastAsia'))
                        en = rFonts.get(qn('w:ascii'))
                    szEl = rPr.find(qn('w:sz'))
                    if szEl is not None:
                        sz = szEl.get(qn('w:val'))
            print(f'  {i}: cn={cn} en={en} sz={sz} text="{text[:40]}"')
print(f'  Total code paragraphs with shading: {code_found}')

print('\n=== Figure captions check ===')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if re.match(r'^图\d+-\d+', text):
        align = para.alignment
        run = para.runs[0] if para.runs else None
        cn = sz = None
        if run:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None: cn = rFonts.get(qn('w:eastAsia'))
                szEl = rPr.find(qn('w:sz'))
                if szEl is not None: sz = szEl.get(qn('w:val'))
        print(f'  {i}: "{text}" align={align} cn={cn} sz={sz}')

print('\n=== Table style check ===')
for i, table in enumerate(doc.tables):
    style = table.style.name if table.style else 'None'
    print(f'  Table {i}: style="{style}"')
