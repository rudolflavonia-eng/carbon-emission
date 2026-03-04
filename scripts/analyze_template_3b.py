"""Round 3b: Analyze tables and Heading style details in template"""
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'E:\2026next\碳排放\论文1.docx')

# Table styles
print('=== TABLES ===')
for i, table in enumerate(doc.tables):
    style = table.style.name if table.style else 'None'
    rows = len(table.rows)
    cols = len(table.columns)
    # First cell text
    cell_text = table.cell(0,0).text.strip()[:30] if rows > 0 and cols > 0 else ''
    print(f'Table {i}: style="{style}" rows={rows} cols={cols} first_cell="{cell_text}"')
    
    # Check table XML for borders
    tPr = table._tbl.find(qn('w:tblPr'))
    if tPr is not None:
        tblBorders = tPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            top = tblBorders.find(qn('w:top'))
            if top is not None:
                print(f'  Border top: val={top.get(qn("w:val"))} sz={top.get(qn("w:sz"))}')
    
    if i >= 5: break

# Heading 1 style details - check actual paragraphs 
print('\n=== HEADING STYLE DETAILS FROM ACTUAL PARAGRAPHS ===')
for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else 'None'
    text = para.text.strip()
    if style.startswith('Heading') and text:
        pPr = para._element.find(qn('w:pPr'))
        
        # Get pBdr
        border_info = ''
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                border_info = 'has_border '
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                border_info += f'jc={jc.get(qn("w:val"))} '
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                border_info += f'sp(line={spacing.get(qn("w:line"))},before={spacing.get(qn("w:before"))},after={spacing.get(qn("w:after"))}) '
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                border_info += f'ind(fl={ind.get(qn("w:firstLine"))}) '
        
        # Run details
        for r in para.runs[:3]:
            rPr = r._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                szEl = rPr.find(qn('w:sz'))
                bEl = rPr.find(qn('w:b'))
                cn = rFonts.get(qn('w:eastAsia')) if rFonts is not None else None
                en = rFonts.get(qn('w:ascii')) if rFonts is not None else None
                sz = szEl.get(qn('w:val')) if szEl is not None else None
                bold = 'true' if bEl is not None and bEl.get(qn('w:val')) != '0' else None
                border_info += f'run(cn={cn},en={en},sz={sz},b={bold}) '
        
        print(f'{i}: [{style}] "{text[:40]}" {border_info}')
        if i > 200:
            break

# Normal style rPr defaults
print('\n=== NORMAL STYLE DEFAULT rPr ===')
normal_style = doc.styles['Normal']
rPr = normal_style.element.find(qn('w:rPr'))
if rPr is not None:
    from lxml import etree
    print(etree.tostring(rPr, encoding='unicode', pretty_print=True))
pPr = normal_style.element.find(qn('w:pPr'))
if pPr is not None:
    from lxml import etree
    print(etree.tostring(pPr, encoding='unicode', pretty_print=True))

# Heading 1 style rPr
print('\n=== HEADING 1 STYLE DEFAULT ===')
h1_style = doc.styles['Heading 1']
rPr = h1_style.element.find(qn('w:rPr'))
if rPr is not None:
    from lxml import etree
    print(etree.tostring(rPr, encoding='unicode', pretty_print=True))
pPr = h1_style.element.find(qn('w:pPr'))
if pPr is not None:
    from lxml import etree
    print(etree.tostring(pPr, encoding='unicode', pretty_print=True))

# Heading 2 style
print('\n=== HEADING 2 STYLE DEFAULT ===')
h2_style = doc.styles['Heading 2']
rPr = h2_style.element.find(qn('w:rPr'))
if rPr is not None:
    from lxml import etree
    print(etree.tostring(rPr, encoding='unicode', pretty_print=True))
pPr = h2_style.element.find(qn('w:pPr'))
if pPr is not None:
    from lxml import etree
    print(etree.tostring(pPr, encoding='unicode', pretty_print=True))

# Heading 3 style
print('\n=== HEADING 3 STYLE DEFAULT ===')
h3_style = doc.styles['Heading 3']
rPr = h3_style.element.find(qn('w:rPr'))
if rPr is not None:
    from lxml import etree
    print(etree.tostring(rPr, encoding='unicode', pretty_print=True))
pPr = h3_style.element.find(qn('w:pPr'))
if pPr is not None:
    from lxml import etree
    print(etree.tostring(pPr, encoding='unicode', pretty_print=True))
