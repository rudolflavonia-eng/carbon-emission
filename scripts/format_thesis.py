"""
Complete thesis formatting script - executes all steps sequentially.
Based on template analysis results.

Template format summary:
- Normal: 宋体/TNR, 小四(12pt/sz24), 两端对齐, 固定值20磅行距, 首行缩进2字符
- Heading 1: 黑体/TNR, 三号(16pt/sz32), 居中, 段前1行段后1行, 无缩进
- Heading 2: 宋体/TNR, 小三(14pt/sz28), 加粗, 左对齐, 段前1行段后1行, 无缩进
- Heading 3: 宋体/TNR, 小四(12pt/sz24), 加粗, 两端对齐, 段前0.5行段后0.5行, 无缩进
- Fig/Tab: 楷体, 五号(10.5pt/sz21), 居中, 无缩进
- Ref item: 宋体/TNR, 五号(10.5pt), List Paragraph style
- Code: Consolas, 小五(9pt), 单倍行距, 灰色底纹+边框
- Page: A4, margins top/bottom=2.5cm, left=3.5cm, right=2.5cm
"""

import re, pickle, sys
sys.path.insert(0, r'E:\2026next\碳排放\scripts')
from thesis_utils import *
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCPATH = r'E:\2026next\碳排放\基于Django的碳排放分析及预测可视化平台的设计与实现_排版后.docx'

# Load classifications
with open(r'E:\2026next\碳排放\scripts\classifications.pkl', 'rb') as f:
    classifications = pickle.load(f)
cls_map = {i: cls for i, cls, _ in classifications}
text_map = {i: text for i, _, text in classifications}

doc = Document(DOCPATH)
total = len(doc.paragraphs)
print(f'Total paragraphs: {total}')

# ============================================================
# STEP 1: Page margins
# ============================================================
print('\n=== STEP 1: Page margins ===')
for i, sec in enumerate(doc.sections):
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    if i == 0:
        # Cover section might have different left margin
        sec.left_margin = Cm(2.5)
    else:
        sec.left_margin = Cm(3.5)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.75)
print(f'  Set margins for {len(doc.sections)} sections.')

# Since this doc likely has 1 section, set it to 3.5cm left
if len(doc.sections) == 1:
    doc.sections[0].left_margin = Cm(3.5)

doc.save(DOCPATH)
print('  Saved.')

# ============================================================
# STEP 2: Set Normal style defaults
# ============================================================
print('\n=== STEP 2: Normal style defaults ===')
doc = Document(DOCPATH)

normal = doc.styles['Normal']
# pPr
pPr = normal.element.find(qn('w:pPr'))
if pPr is None:
    pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
    normal.element.append(pPr)

# Spacing: fixed 20pt (400 twips)
spacing = pPr.find(qn('w:spacing'))
if spacing is None:
    spacing = parse_xml(f'<w:spacing {nsdecls("w")}/>')
    pPr.append(spacing)
spacing.set(qn('w:line'), '400')
spacing.set(qn('w:lineRule'), 'exact')

# Indent: firstLine 200 chars (2 chars)
ind = pPr.find(qn('w:ind'))
if ind is None:
    ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
    pPr.append(ind)
ind.set(qn('w:firstLine'), '200')
ind.set(qn('w:firstLineChars'), '200')

# Alignment: both
jc = pPr.find(qn('w:jc'))
if jc is None:
    jc = parse_xml(f'<w:jc {nsdecls("w")}/>')
    pPr.append(jc)
jc.set(qn('w:val'), 'both')

# rPr: 宋体/TNR, 小四
rPr = normal.element.find(qn('w:rPr'))
if rPr is None:
    rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
    normal.element.append(rPr)
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:eastAsia'), '宋体')
sz = rPr.find(qn('w:sz'))
if sz is None:
    sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
    rPr.append(sz)
sz.set(qn('w:val'), '24')  # 小四

doc.save(DOCPATH)
print('  Saved Normal style.')

# ============================================================
# STEP 3: Set Heading styles
# ============================================================  
print('\n=== STEP 3: Heading styles ===')
doc = Document(DOCPATH)

# --- Heading 1: 黑体, 三号(16pt), 居中, 段前1行段后1行 ---
h1_style = doc.styles['Heading 1']
h1_pPr = h1_style.element.find(qn('w:pPr'))
if h1_pPr is None:
    h1_pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
    h1_style.element.append(h1_pPr)

# Center alignment
h1_jc = h1_pPr.find(qn('w:jc'))
if h1_jc is None:
    h1_jc = parse_xml(f'<w:jc {nsdecls("w")}/>')
    h1_pPr.append(h1_jc)
h1_jc.set(qn('w:val'), 'center')

# Spacing: before=1line (100), after=1line (100)
h1_sp = h1_pPr.find(qn('w:spacing'))
if h1_sp is None:
    h1_sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
    h1_pPr.append(h1_sp)
h1_sp.set(qn('w:beforeLines'), '100')
h1_sp.set(qn('w:before'), '312')
h1_sp.set(qn('w:afterLines'), '100')
h1_sp.set(qn('w:after'), '312')

# No indent
h1_ind = h1_pPr.find(qn('w:ind'))
if h1_ind is None:
    h1_ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
    h1_pPr.append(h1_ind)
h1_ind.set(qn('w:firstLine'), '0')
h1_ind.set(qn('w:firstLineChars'), '0')

# Outline level
h1_ol = h1_pPr.find(qn('w:outlineLvl'))
if h1_ol is None:
    h1_ol = parse_xml(f'<w:outlineLvl {nsdecls("w")}/>')
    h1_pPr.append(h1_ol)
h1_ol.set(qn('w:val'), '0')

# rPr: 黑体, 三号 16pt
h1_rPr = h1_style.element.find(qn('w:rPr'))
if h1_rPr is None:
    h1_rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
    h1_style.element.append(h1_rPr)
h1_rFonts = h1_rPr.find(qn('w:rFonts'))
if h1_rFonts is None:
    h1_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
    h1_rPr.insert(0, h1_rFonts)
h1_rFonts.set(qn('w:eastAsia'), '黑体')
h1_sz = h1_rPr.find(qn('w:sz'))
if h1_sz is None:
    h1_sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
    h1_rPr.append(h1_sz)
h1_sz.set(qn('w:val'), '32')
# Remove bold if exists (template has bCs but not b)
h1_b = h1_rPr.find(qn('w:b'))
if h1_b is not None:
    h1_rPr.remove(h1_b)

# --- Heading 2: 宋体/TNR, 小三(14pt), 加粗, 左对齐, 段前1行段后1行 ---
h2_style = doc.styles['Heading 2']
h2_pPr = h2_style.element.find(qn('w:pPr'))
if h2_pPr is None:
    h2_pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
    h2_style.element.append(h2_pPr)

h2_jc = h2_pPr.find(qn('w:jc'))
if h2_jc is None:
    h2_jc = parse_xml(f'<w:jc {nsdecls("w")}/>')
    h2_pPr.append(h2_jc)
h2_jc.set(qn('w:val'), 'left')

h2_sp = h2_pPr.find(qn('w:spacing'))
if h2_sp is None:
    h2_sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
    h2_pPr.append(h2_sp)
h2_sp.set(qn('w:beforeLines'), '100')
h2_sp.set(qn('w:before'), '312')
h2_sp.set(qn('w:afterLines'), '100')
h2_sp.set(qn('w:after'), '312')

h2_ind = h2_pPr.find(qn('w:ind'))
if h2_ind is None:
    h2_ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
    h2_pPr.append(h2_ind)
h2_ind.set(qn('w:firstLine'), '0')
h2_ind.set(qn('w:firstLineChars'), '0')

h2_ol = h2_pPr.find(qn('w:outlineLvl'))
if h2_ol is None:
    h2_ol = parse_xml(f'<w:outlineLvl {nsdecls("w")}/>')
    h2_pPr.append(h2_ol)
h2_ol.set(qn('w:val'), '1')

h2_rPr = h2_style.element.find(qn('w:rPr'))
if h2_rPr is None:
    h2_rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
    h2_style.element.append(h2_rPr)
h2_b = h2_rPr.find(qn('w:b'))
if h2_b is None:
    h2_b = parse_xml(f'<w:b {nsdecls("w")}/>')
    h2_rPr.append(h2_b)
h2_sz = h2_rPr.find(qn('w:sz'))
if h2_sz is None:
    h2_sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
    h2_rPr.append(h2_sz)
h2_sz.set(qn('w:val'), '28')

# --- Heading 3: 宋体/TNR, 小四(12pt), 加粗, 段前0.5行段后0.5行 ---
h3_style = doc.styles['Heading 3']
h3_pPr = h3_style.element.find(qn('w:pPr'))
if h3_pPr is None:
    h3_pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
    h3_style.element.append(h3_pPr)

h3_sp = h3_pPr.find(qn('w:spacing'))
if h3_sp is None:
    h3_sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
    h3_pPr.append(h3_sp)
h3_sp.set(qn('w:beforeLines'), '50')
h3_sp.set(qn('w:before'), '156')
h3_sp.set(qn('w:afterLines'), '50')
h3_sp.set(qn('w:after'), '156')

h3_ind = h3_pPr.find(qn('w:ind'))
if h3_ind is None:
    h3_ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
    h3_pPr.append(h3_ind)
h3_ind.set(qn('w:firstLine'), '0')
h3_ind.set(qn('w:firstLineChars'), '0')

h3_ol = h3_pPr.find(qn('w:outlineLvl'))
if h3_ol is None:
    h3_ol = parse_xml(f'<w:outlineLvl {nsdecls("w")}/>')
    h3_pPr.append(h3_ol)
h3_ol.set(qn('w:val'), '2')

h3_rPr = h3_style.element.find(qn('w:rPr'))
if h3_rPr is None:
    h3_rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
    h3_style.element.append(h3_rPr)
h3_b = h3_rPr.find(qn('w:b'))
if h3_b is None:
    h3_b = parse_xml(f'<w:b {nsdecls("w")}/>')
    h3_rPr.append(h3_b)

doc.save(DOCPATH)
print('  Saved heading styles.')

# ============================================================
# STEP 4: Apply styles to paragraphs and format runs
# ============================================================
print('\n=== STEP 4: Apply styles and format all paragraphs ===')
doc = Document(DOCPATH)

h1_count = h2_count = h3_count = body_count = fig_count = tab_count = 0
code_count = code_marker_count = ref_count = 0

for i, para in enumerate(doc.paragraphs):
    cls = cls_map.get(i, 'body')
    
    if cls == 'paper_title':
        # Title: 黑体, 二号(22pt), 居中, 加粗
        para.style = doc.styles['Normal']
        set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
        clear_indent(para)
        fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=22, b=True)
    
    elif cls == 'abs_cn_title':
        # 摘要标题: same as Heading 1
        para.style = doc.styles['Heading 1']
        fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=16, b=False)
        clear_indent(para)
    
    elif cls == 'abs_en_title':
        para.style = doc.styles['Heading 1']
        fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=16, b=False)
        clear_indent(para)
    
    elif cls == 'abs_cn':
        # 中文摘要正文: same as normal body
        para.style = doc.styles['Normal']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=False)
    
    elif cls == 'abs_en':
        para.style = doc.styles['Normal']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=False)
    
    elif cls == 'kw_cn':
        # Keywords: "关键词" bold, rest not bold
        para.style = doc.styles['Normal']
        for run in para.runs:
            if '关键词' in run.text or '关 键 词' in run.text:
                set_run_font(run, cn='宋体', en='Times New Roman', sz=12, b=True)
            else:
                set_run_font(run, cn='宋体', en='Times New Roman', sz=12, b=False)
    
    elif cls == 'kw_en':
        para.style = doc.styles['Normal']
        for run in para.runs:
            txt = run.text.lower()
            if 'keyword' in txt or 'key word' in txt:
                set_run_font(run, cn='宋体', en='Times New Roman', sz=12, b=True)
            else:
                set_run_font(run, cn='宋体', en='Times New Roman', sz=12, b=False)
    
    elif cls == 'toc_title':
        # 目录标题: same as h1
        para.style = doc.styles['Heading 1']
        fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=16, b=False)
        clear_indent(para)
    
    elif cls == 'toc':
        # Don't touch TOC entries - they'll be regenerated
        pass
    
    elif cls == 'h1':
        para.style = doc.styles['Heading 1']
        fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=16, b=False)
        clear_indent(para)
        h1_count += 1
    
    elif cls == 'h2':
        para.style = doc.styles['Heading 2']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=14, b=True)
        clear_indent(para)
        h2_count += 1
    
    elif cls == 'h3':
        para.style = doc.styles['Heading 3']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=True)
        clear_indent(para)
        h3_count += 1
    
    elif cls == 'fig':
        # Figure captions: 楷体, 五号(10.5pt), 居中
        # First strip <center> tags from text
        para.style = doc.styles['Normal']
        origin_text = text_map.get(i, '')
        # Clear existing runs and set clean text
        if para.runs:
            # Set all runs to caption format
            for run in para.runs:
                set_run_font(run, cn='楷体', en='楷体', sz=10.5, b=False)
            # Replace the <center>...</center> text
            full_text = para.text
            if '<center>' in full_text:
                clean = re.sub(r'</?center>', '', full_text)
                # Clear all runs except first
                for j in range(len(para.runs)-1, 0, -1):
                    para.runs[j].text = ''
                para.runs[0].text = clean
        set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
        clear_indent(para)
        # Set single line spacing for caption
        pPr = ensure_ppr(para)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:line'), '400')
        sp.set(qn('w:lineRule'), 'exact')
        fig_count += 1
    
    elif cls == 'tab':
        # Table captions: same as figure
        para.style = doc.styles['Normal']
        if para.runs:
            for run in para.runs:
                set_run_font(run, cn='楷体', en='楷体', sz=10.5, b=False)
            full_text = para.text
            if '<center>' in full_text:
                clean = re.sub(r'</?center>', '', full_text)
                for j in range(len(para.runs)-1, 0, -1):
                    para.runs[j].text = ''
                para.runs[0].text = clean
        set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
        clear_indent(para)
        pPr = ensure_ppr(para)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:line'), '400')
        sp.set(qn('w:lineRule'), 'exact')
        tab_count += 1
    
    elif cls == 'img_ref':
        # Image references: center, no indent
        para.style = doc.styles['Normal']
        set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
        clear_indent(para)
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=False)
    
    elif cls == 'code':
        # Code: Consolas, 小五(9pt), single line spacing, gray bg + border
        para.style = doc.styles['Normal']
        fmt_all_runs(para, cn='Consolas', en='Consolas', sz=9, b=False)
        set_paragraph_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT)
        clear_indent(para)
        add_code_border(para)
        # Single line spacing (240 twips = ~12pt)
        pPr = ensure_ppr(para)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:line'), '240')
        sp.set(qn('w:lineRule'), 'exact')
        # Remove space before/after for compact code
        for attr in [qn('w:before'), qn('w:after'), qn('w:beforeLines'), qn('w:afterLines')]:
            if sp.get(attr):
                del sp.attrib[attr]
        code_count += 1
    
    elif cls == 'code_marker':
        # Hide code markers (```python, ```)
        para.style = doc.styles['Normal']
        # Clear text
        for run in para.runs:
            run.text = ''
        # Minimal height
        clear_indent(para)
        pPr = ensure_ppr(para)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:line'), '20')
        sp.set(qn('w:lineRule'), 'exact')
        for attr in [qn('w:before'), qn('w:after'), qn('w:beforeLines'), qn('w:afterLines')]:
            if sp.get(attr):
                del sp.attrib[attr]
        code_marker_count += 1
    
    elif cls == 'ref_title':
        # Reference title: same as h1
        para.style = doc.styles['Heading 1']
        fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=16, b=False)
        clear_indent(para)
    
    elif cls == 'ref_item':
        # Reference items: 宋体/TNR, 五号(10.5pt)
        para.style = doc.styles['Normal']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=10.5, b=False)
        set_hanging_indent(para, hanging=420, left=420)
        # Keep normal line spacing
        pPr = ensure_ppr(para)
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
            pPr.append(sp)
        sp.set(qn('w:line'), '400')
        sp.set(qn('w:lineRule'), 'exact')
        ref_count += 1
    
    elif cls == 'ack_title':
        para.style = doc.styles['Heading 1']
        fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=16, b=False)
        clear_indent(para)
    
    elif cls == 'ack_body':
        para.style = doc.styles['Normal']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=False)
    
    elif cls == 'appendix_title':
        para.style = doc.styles['Heading 1']
        fmt_all_runs(para, cn='黑体', en='Times New Roman', sz=16, b=False)
        clear_indent(para)
    
    elif cls == 'appendix_section':
        para.style = doc.styles['Heading 2']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=14, b=True)
        clear_indent(para)
    
    elif cls == 'appendix_body':
        para.style = doc.styles['Normal']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=False)
    
    elif cls == 'body':
        # Normal body text
        para.style = doc.styles['Normal']
        fmt_all_runs(para, cn='宋体', en='Times New Roman', sz=12, b=False)
        body_count += 1
    
    elif cls == 'empty':
        pass

print(f'  h1={h1_count} h2={h2_count} h3={h3_count} body={body_count}')
print(f'  fig={fig_count} tab={tab_count} code={code_count} code_marker={code_marker_count} ref={ref_count}')

doc.save(DOCPATH)
print('  Saved paragraph formatting.')

# ============================================================
# STEP 5: Set tables to "表格" style (Table Grid)
# ============================================================
print('\n=== STEP 5: Format tables ===')
doc = Document(DOCPATH)

# Check available table styles
table_style_names = [s.name for s in doc.styles if s.type is not None and s.type.name == 'TABLE']
print(f'  Available table styles: {table_style_names[:10]}')

for i, table in enumerate(doc.tables):
    # Apply Table Grid style (closest to "表格")
    try:
        table.style = 'Table Grid'
    except:
        try:
            table.style = 'TableGrid'
        except:
            print(f'  Warning: Could not set table style for table {i}')
    
    # Format table cell text
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, cn='宋体', en='Times New Roman', sz=10.5, b=False)

print(f'  Formatted {len(doc.tables)} tables.')
doc.save(DOCPATH)
print('  Saved.')

# ============================================================
# STEP 6: Add line break before figure/table captions
# User requirement: 题注前要加换行符
# ============================================================
print('\n=== STEP 6: Add line break before captions ===')
doc = Document(DOCPATH)

caption_indices = set()
for i, cls, _ in classifications:
    if cls in ('fig', 'tab'):
        caption_indices.add(i)

# For each caption, add an empty paragraph before it (line break)
# We need to insert in reverse order to maintain indices
from lxml import etree

sorted_caption_indices = sorted(caption_indices, reverse=True)
inserted = 0
for idx in sorted_caption_indices:
    para = doc.paragraphs[idx]
    # Insert an empty paragraph before this one
    new_p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:line="100" w:lineRule="exact"/><w:ind w:firstLine="0" w:firstLineChars="0"/></w:pPr></w:p>')
    para._element.addprevious(new_p)
    inserted += 1

print(f'  Inserted {inserted} line breaks before captions.')
doc.save(DOCPATH)
print('  Saved.')

# ============================================================
# STEP 7: Chapter page breaks (h1)
# ============================================================
print('\n=== STEP 7: Chapter page breaks ===')
doc = Document(DOCPATH)

# Re-find h1 paragraphs after insertions
h1_found = 0
first_h1 = True
for para in doc.paragraphs:
    text = para.text.strip()
    if re.match(r'^第\d+章\s', text) or text in ['摘  要', '摘要', 'Abstract', 'ABSTRACT', '参考文献', '致  谢', '致谢', '附  录', '附录']:
        if first_h1:
            first_h1 = False
            # Skip first h1
        else:
            add_page_break_before(para)
            h1_found += 1

print(f'  Added page breaks before {h1_found} chapter headings.')
doc.save(DOCPATH)
print('  Saved.')

# ============================================================
# STEP 8: Verification
# ============================================================
print('\n=== STEP 8: Verification ===')
doc = Document(DOCPATH)

# Sample check
checks = {'h1': 0, 'h2': 0, 'h3': 0, 'body': 0, 'fig': 0, 'tab': 0, 'code': 0}
for para in doc.paragraphs:
    style = para.style.name
    if style == 'Heading 1':
        checks['h1'] += 1
    elif style == 'Heading 2':
        checks['h2'] += 1
    elif style == 'Heading 3':
        checks['h3'] += 1

print(f'  Style counts: Heading1={checks["h1"]}, Heading2={checks["h2"]}, Heading3={checks["h3"]}')
print(f'  Total paragraphs: {len(doc.paragraphs)}')
print(f'  Total tables: {len(doc.tables)}')

# Check a sample body paragraph
for para in doc.paragraphs:
    text = para.text.strip()
    if len(text) > 50 and para.style.name == 'Normal':
        has_run = len(para.runs) > 0
        if has_run:
            r = para.runs[0]
            rPr = r._element.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                cn = rFonts.get(qn('w:eastAsia')) if rFonts is not None else None
                en = rFonts.get(qn('w:ascii')) if rFonts is not None else None
                szEl = rPr.find(qn('w:sz'))
                sz_val = szEl.get(qn('w:val')) if szEl is not None else None
                print(f'  Sample body: cn={cn}, en={en}, sz={sz_val}')
                break

print('\n=== FORMATTING COMPLETE ===')
print(f'Output: {DOCPATH}')
